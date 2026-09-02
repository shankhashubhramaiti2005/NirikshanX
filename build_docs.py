import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(26, 35, 126) # Deep Navy
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 110, 247) # Accent Blue
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
    return h

def add_body_paragraph(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(30, 41, 59)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_item(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(30, 41, 59)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout_box(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"📌 {title}: ")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(10.5)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(79, 110, 247)
    
    r_txt = p.add_run(text)
    r_txt.font.name = 'Calibri'
    r_txt.font.size = Pt(10.5)
    r_txt.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def format_table_headers(table, headers, col_widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1A237E") # Deep Navy
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        if col_widths and i < len(col_widths):
            hdr_cells[i].width = col_widths[i]

def style_table_rows(table, data, col_widths=None):
    for r_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_hex = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_hex)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(30, 41, 59)
            if col_widths and c_idx < len(col_widths):
                row_cells[c_idx].width = col_widths[c_idx]

def create_project_documentation():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run("NirikshanX — Technical Documentation & Interview Master Guide")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(24)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(26, 35, 126)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("AI-Powered Packaged Commodity Compliance Inspection System\nLegal Metrology (Packaged Commodities) Rules, 2011 & Legal Metrology Act, 2009")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    add_callout_box(doc, "This document contains complete system design, architecture, AI pipeline specifications, database models, API contracts, solved edge cases, and 15 technical interview Q&A questions for NirikshanX.", "EXECUTIVE BRIEF")

    # 1. SYSTEM OVERVIEW & PROBLEM STATEMENT
    add_heading_1(doc, "1. System Overview & Problem Statement")
    add_body_paragraph(doc, "NirikshanX is an enterprise-grade, AI-driven regulatory compliance inspection platform designed for enforcement officers, legal metrology inspectors, and consumer protection authorities in India. It automates the verification of mandatory regulatory declarations on packaged commodities under the Legal Metrology (Packaged Commodities) Rules, 2011.")
    
    add_heading_2(doc, "The Regulatory & Business Problem")
    add_bullet_item(doc, "Millions of packaged goods (food, cosmetics, electronics, pharmaceuticals) are distributed daily across Indian retail channels.", "Scale of Enforcement: ")
    add_bullet_item(doc, "Manual inspection of packaging labels for mandatory declarations (MRP, Net Qty, Mfd Date, Manufacturer details, FSSAI Lic) is slow, expensive, and subject to human oversight.", "Manual Inspection Bottleneck: ")
    add_bullet_item(doc, "Naïve OCR systems often evaluate compliance purely as (violations == 0 -> COMPLIANT). When a user uploads a non-product image (person, animal, landscape, screenshot, blank paper), traditional systems fail to detect violations and falsely output 100% compliance.", "The 100% False Positive AI Bug: ")
    add_bullet_item(doc, "NirikshanX solves this by introducing a Multi-Stage AI Validation Pipeline that verifies image quality, detects packaged commodities, extracts OCR text using ONNX deep learning, and enforces positive declaration evidence before assigning compliance scores.", "The NirikshanX Solution: ")

    # 2. MULTI-STAGE AI VALIDATION PIPELINE
    add_heading_1(doc, "2. Multi-Stage AI Validation Pipeline Architecture")
    add_body_paragraph(doc, "NirikshanX replaces blind compliance scoring with a 6-stage sequential AI pipeline. Every uploaded image passes through strict quality and object detection gates before OCR or rule engine evaluation occurs.")

    headers_pipeline = ["Stage", "Component", "Implementation", "Decision Output"]
    widths_pipeline = [Inches(1.0), Inches(1.8), Inches(2.2), Inches(1.5)]
    tbl_pipe = doc.add_table(rows=1, cols=4)
    format_table_headers(tbl_pipe, headers_pipeline, widths_pipeline)

    data_pipeline = [
        ["Stage 1", "Image Quality Check", "quality_service.py\n(Resolution, Blur, Exposure, StdDev)", "If unusable -> IMAGE_QUALITY_INSUFFICIENT"],
        ["Stage 2", "Packaged Product Detection", "product_detection_service.py\n(Skin, Landscape, Document, Edge Grid)", "If confidence < 0.60 -> NOT_A_PRODUCT"],
        ["Stage 3", "Text Region Detection & OCR", "ocr_service.py\n(RapidOCR ONNX Deep Learning Model)", "Text bounding boxes + raw text extraction"],
        ["Stage 4", "Declaration Field Parsing", "ocr_service.py\n(Regex + NLP Pattern Extraction)", "Parsed dict of 15 Legal Metrology fields"],
        ["Stage 5", "Legal Metrology Rule Engine", "engine.py\n(LM (PC) Rules 2011 Evaluation)", "Rule check results (PASS / FAIL / WARN)"],
        ["Stage 6", "Positive Evidence Enforcement", "pipeline.py\n(Evidence Count & Inspection Conf)", "COMPLIANT / NON_COMPLIANT / NEEDS_REVIEW"],
    ]
    style_table_rows(tbl_pipe, data_pipeline, widths_pipeline)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2(doc, "Deep Dive into Pipeline Stages")
    add_bullet_item(doc, "Calculates pixel resolution, mean luminance (detecting dark < 15.0 or overexposed > 240.0 images), color standard deviation (detecting blank solid images < 8.0), and Laplacian edge magnitude (detecting blur < 10.0).", "1. Image Quality Validation: ")
    add_bullet_item(doc, "Evaluates human skin-tone ratio (detecting person/face photos > 0.20), nature sky/vegetation ratio (detecting outdoor landscapes > 0.25), document paper ratio (detecting flat text scans > 0.75), and high-contrast packaging declaration grid cells. Threshold set to PRODUCT_DETECTION_THRESHOLD = 0.60.", "2. Packaged Product / Commodity Detection: ")
    add_bullet_item(doc, "Employs RapidOCR (ONNX Runtime) deep learning models for text detection and text recognition. Operates standalone on CPU/GPU without external dependencies like Tesseract EXE.", "3. RapidOCR Deep Learning Engine: ")
    add_bullet_item(doc, "Extracts 15 regulatory fields using regex pattern matching and contextual token parsing (MRP, Net Qty, FSSAI 14-digit, Mfd Date, Best Before, Manufacturer Name & Address, Country of Origin, Ingredients, Nutritional Info, Allergen Info, Batch No, BIS Lic, Customer Care).", "4. Legal Metrology Field Parser: ")
    add_bullet_item(doc, "If extracted evidence count < 2 fields or overall inspection confidence < 60%, the system conservatively forces the status to NEEDS_REVIEW (Human-in-the-loop inspection) rather than claiming fake compliance.", "5. Positive Evidence Rule Enforcement: ")

    # 3. COMPLETE TECHNOLOGY STACK
    add_heading_1(doc, "3. Complete Technology Stack & System Dependencies")
    
    headers_tech = ["Layer", "Technology / Library", "Version", "Purpose & Role in Project"]
    widths_tech = [Inches(1.2), Inches(1.8), Inches(1.0), Inches(2.5)]
    tbl_tech = doc.add_table(rows=1, cols=4)
    format_table_headers(tbl_tech, headers_tech, widths_tech)

    data_tech = [
        ["Backend Core", "FastAPI", "0.111.0", "Asynchronous RESTful API framework"],
        ["ASGI Server", "Uvicorn", "0.29.0", "High-performance Python web server"],
        ["Database ORM", "SQLAlchemy", "2.0.30", "Object-Relational Mapping & Database Layer"],
        ["Database", "SQLite / PostgreSQL", "3.x", "Relational persistent scan & case storage"],
        ["AI / OCR Engine", "RapidOCR (ONNX)", "1.2.3", "Deep learning text detection & recognition"],
        ["Computer Vision", "Pillow (PIL) & OpenCV", "12.3.0 / 5.0", "Image preprocessing, quality analytics, feature ratios"],
        ["Authentication", "PyJWT & Passlib", "3.3.0 / 1.7", "JWT OAuth2 authentication & bcrypt hashing"],
        ["PDF Reports", "ReportLab", "4.2.0", "Instant PDF compliance audit report generation"],
        ["Doc Generator", "python-docx", "1.2.0", "Automated Microsoft Word documentation generation"],
        ["Frontend SPA", "HTML5, CSS3, Vanilla JS", "ES6+", "High-performance, zero-dependency SPA with Chart.js"],
    ]
    style_table_rows(tbl_tech, data_tech, widths_tech)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. DATABASE MODELS & SCHEMA DESIGN
    add_heading_1(doc, "4. Database Models & Schema Design")
    add_body_paragraph(doc, "The database schema is designed in SQLAlchemy 2.0 with full relational integrity, cascading deletions, and JSON metadata support.")

    add_heading_2(doc, "Core Database Entities")
    add_bullet_item(doc, "Stores email, full_name, hashed_password, role (ADMIN, INSPECTOR, ENFORCEMENT_OFFICER, USER), and timestamps.", "1. User Entity: ")
    add_bullet_item(doc, "Stores scan ID, product_name, category (FOOD, GENERAL, COSMETICS, ELECTRONICS, MEDICINE), image_path, status (PENDING, PROCESSING, COMPLETED, FAILED, REVIEW_REQUIRED, NOT_A_PRODUCT, IMAGE_QUALITY_INSUFFICIENT), compliance_score (0-100%), overall_status, reviewer_id, notes, and ai_debug_metrics JSON.", "2. Scan Entity: ")
    add_bullet_item(doc, "Stores scan_id, field_name, extracted_value, confidence (0.0-1.0), bounding_box, and is_present boolean.", "3. Declaration Entity: ")
    add_bullet_item(doc, "Stores scan_id, rule_id, field, status (PASS, FAIL, WARNING, REVIEW_REQUIRED, NOT_A_PRODUCT, IMAGE_QUALITY_INSUFFICIENT), severity (LOW, MEDIUM, HIGH, CRITICAL), message, evidence JSON, and confidence.", "4. Violation Entity: ")
    add_bullet_item(doc, "Stores legal enforcement case number (e.g. NX-20260902-0010), scan_id, status (OPEN, UNDER_INVESTIGATION, CLOSED), assigned inspector ID, and notes.", "5. Case Entity: ")

    # 5. LEGAL METROLOGY RULES ENFORCED
    add_heading_1(doc, "5. Legal Metrology (Packaged Commodities) Rules Matrix")
    
    headers_rules = ["Rule ID", "Field Name", "Legal Metrology Reference", "Severity", "Requirement Description"]
    widths_rules = [Inches(1.1), Inches(1.4), Inches(1.3), Inches(0.9), Inches(1.8)]
    tbl_rules = doc.add_table(rows=1, cols=5)
    format_table_headers(tbl_rules, headers_rules, widths_rules)

    data_rules = [
        ["LM-NAME", "product_name", "Rule 4(a)", "HIGH", "Generic / trade name of the commodity"],
        ["LM-NET_QTY", "net_quantity", "Rule 4(b)", "HIGH", "Net weight/volume in standard metric units (g/kg/ml/l)"],
        ["LM-MFR_NAME", "manufacturer_name", "Rule 4(c)", "CRITICAL", "Name of manufacturer, packer, or importer"],
        ["LM-MFR_ADDR", "manufacturer_address", "Rule 4(c)", "CRITICAL", "Complete address including city, state, and PIN code"],
        ["LM-COUNTRY", "country_of_origin", "Rule 4(d)", "MEDIUM", "Country of origin for imported or domestic goods"],
        ["LM-MRP", "mrp", "Rule 4(e)", "CRITICAL", "Maximum Retail Price inclusive of all taxes"],
        ["LM-MFD_DATE", "month_year_of_manufacture", "Rule 4(f)", "HIGH", "Month and year of manufacture or packing"],
        ["LM-BEST_BEFORE", "best_before", "Rule 4(g)", "HIGH", "Best before date / use by date"],
        ["LM-FSSAI", "fssai_license", "FSS Act 2006", "HIGH", "14-digit valid FSSAI license number"],
        ["LM-INGREDIENTS", "ingredients", "Rule 4(h)", "HIGH", "Complete ingredient declaration list"],
        ["LM-NUTRITION", "nutritional_info", "Rule 4(i)", "MEDIUM", "Nutritional information per 100g/100ml"],
        ["LM-ALLERGEN", "allergen_info", "Rule 4(j)", "MEDIUM", "Allergen warning declaration"],
    ]
    style_table_rows(tbl_rules, data_rules, widths_rules)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 6. API ENDPOINTS REFERENCE
    add_heading_1(doc, "6. RESTful API Endpoints & Contract Reference")
    
    headers_api = ["Method", "Endpoint", "Auth Required", "Description & Response Payload"]
    widths_api = [Inches(0.9), Inches(1.8), Inches(1.1), Inches(2.7)]
    tbl_api = doc.add_table(rows=1, cols=4)
    format_table_headers(tbl_api, headers_api, widths_api)

    data_api = [
        ["POST", "/auth/login", "No", "Authenticates user and returns JWT access_token + user profile"],
        ["POST", "/auth/register", "No", "Registers new inspector/admin user and returns JWT token"],
        ["GET", "/auth/me", "Yes", "Returns current authenticated user details"],
        ["POST", "/scans/submit", "Yes", "Submits image scan; triggers background AI validation pipeline"],
        ["GET", "/scans/{id}", "Yes", "Fetches scan results, declarations, violations & ai_debug_metrics"],
        ["GET", "/scans/", "Yes", "Lists scan history for current user or admin"],
        ["GET", "/dashboard/stats", "Yes", "Returns compliance rate, violation counts, category breakdown"],
        ["GET", "/reports/{id}/pdf", "Yes", "Generates & downloads formal PDF compliance audit report"],
    ]
    style_table_rows(tbl_api, data_api, widths_api)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 7. CRITICAL BUG FIXES & REFACTORING
    add_heading_1(doc, "7. Critical Architectural Bug Fixes & Refactoring")
    
    add_heading_2(doc, "1. Elimination of 100% False Positive Compliance Bug")
    add_body_paragraph(doc, "Root Cause: The original implementation unconditionally invoked run_mock_pipeline, returning pre-populated static declarations (DEMO_COMPLIANT) for all uploaded images when scenario=0 (default form upload). The rule engine scored this dummy data as 100% compliant.")
    add_body_paragraph(doc, "Fix: Created pipeline.py orchestrator. Standard user uploads now run the real multi-stage AI pipeline. Demo mock data only triggers if DEMO_MODE=true AND scenario > 0 is explicitly selected.")

    add_heading_2(doc, "2. Real Deep Learning OCR Integration")
    add_body_paragraph(doc, "Root Cause: Original code contained no active OCR model, falling back to empty text parsing and returning 0% confidence.")
    add_body_paragraph(doc, "Fix: Installed rapidocr-onnxruntime. Integrated RapidOCR deep learning ONNX models into OCRService. Text bounding boxes and high-confidence declaration strings are extracted directly from label photos.")

    add_heading_2(doc, "3. SQLite Database Migration & Telemetry Field")
    add_body_paragraph(doc, "Root Cause: Adding ai_debug_metrics JSON column to Scan model caused OperationalError: no such column: scans.ai_debug_metrics on existing SQLite database files.")
    add_body_paragraph(doc, "Fix: Executed ALTER TABLE scans ADD COLUMN ai_debug_metrics JSON; and added automatic startup column migration logic in main.py.")

    # 8. TOP 15 INTERVIEW QUESTIONS & ANSWERS
    add_heading_1(doc, "8. Top 15 Technical Interview Questions & Model Answers")

    qas = [
        ("Q1: What is NirikshanX and what problem does it solve?",
         "NirikshanX is an AI-powered compliance inspection system built for Legal Metrology enforcement in India. It automates the verification of mandatory declarations on packaged commodity labels (MRP, Net Quantity, Mfd Date, Manufacturer details, FSSAI license) under the Legal Metrology (Packaged Commodities) Rules, 2011."),

        ("Q2: How did you handle false positives where random non-product images were marked 100% compliant?",
         "We established a 6-stage AI validation pipeline. Before running OCR or compliance rules, the system passes images through Quality Checking (detecting resolution, blur, dark/bright lighting, solid blank images) and Product Detection (detecting skin tone ratios for human photos, sky/vegetation ratios for landscape photos, and plain paper document layouts). If product detection confidence is below 60%, the pipeline immediately halts and returns NOT_A_PRODUCT."),

        ("Q3: Why did you choose RapidOCR over PyTesseract or EasyOCR?",
         "RapidOCR uses ONNX Runtime deep learning models (PP-OCRv4 architecture), which provides superior speed (CPU optimized), high accuracy on stylized/curved packaged label text, and operates standalone without requiring an external C++ executable (like Tesseract.exe) or heavy PyTorch GPU dependencies (like EasyOCR)."),

        ("Q4: How does NirikshanX calculate overall inspection confidence?",
         "Inspection confidence is calculated across four weighted components:\n"
         "Inspection Confidence = (0.20 * Quality Score) + (0.30 * Product Detection Conf) + (0.25 * OCR Conf) + (0.25 * Field Extraction Conf).\n"
         "If extracted evidence is below 2 fields or overall confidence < 60%, the system conservatively outputs NEEDS_REVIEW (Human Inspector Review)."),

        ("Q5: What happens if an image has 0 violations but very low OCR text confidence?",
         "To prevent the 'No Violations = Compliant' bug, NirikshanX enforces Positive Evidence Validation. If OCR returns insufficient readable evidence (fewer than 2 mandatory fields parsed), the system outputs NEEDS_REVIEW instead of COMPLIANT."),

        ("Q6: Describe the tech stack used in NirikshanX.",
         "Backend: FastAPI (Python 3.13), Uvicorn, SQLAlchemy 2.0 ORM, SQLite/PostgreSQL, Pydantic v2.\n"
         "AI/Vision: RapidOCR (ONNX), OpenCV, Pillow (PIL).\n"
         "Auth & Security: OAuth2 JWT tokens with passlib (bcrypt) hashing.\n"
         "Frontend: Vanilla JavaScript SPA, HTML5, CSS3 with custom design tokens, Chart.js.\n"
         "Reports: ReportLab (PDF) and python-docx (Word)."),

        ("Q7: How are legal metrology rules modeled in the backend?",
         "Rules are encapsulated in engine.py. Each rule defines a rule ID (e.g. LM-MRP), field key, severity (CRITICAL, HIGH, MEDIUM, LOW), rule reference (e.g. Rule 4(e)), validation regex check (e.g. 14-digit FSSAI check, MRP currency symbol check), and evidence payload."),

        ("Q8: How do you handle asynchronous scan processing in FastAPI?",
         "When a user submits a scan via POST /scans/submit, FastAPI creates a database record with PENDING status and delegates the execution of run_validation_pipeline to BackgroundTasks. The frontend polls GET /scans/{id} until status transitions to COMPLETED, REVIEW_REQUIRED, NOT_A_PRODUCT, or IMAGE_QUALITY_INSUFFICIENT."),

        ("Q9: What database schema changes were made for AI telemetry?",
         "We added an ai_debug_metrics JSON column to the Scan table. This stores stage-by-stage AI metrics (quality score, product confidence, label confidence, OCR confidence, extraction confidence, overall confidence, model availability, and execution reasoning) rendered in the UI AI Debug Panel."),

        ("Q10: How do you handle database migration when adding new columns in SQLite?",
         "SQLAlchemy's create_all() does not alter existing SQLite tables. In main.py's on_startup hook, we execute an explicit ALTER TABLE scans ADD COLUMN ai_debug_metrics JSON inside a safe try/except block."),

        ("Q11: How is user authentication and authorization implemented?",
         "Authentication uses JWT Bearer tokens signed with HS256 algorithm. Passlib bcrypt hashes passwords. Role-Based Access Control (RBAC) supports ADMIN, INSPECTOR, ENFORCEMENT_OFFICER, and USER roles. Admin/inspectors can view all scans, while standard users can only view their own scans."),

        ("Q12: How does the system distinguish between a human photo and a packaged product?",
         "ProductDetectionService converts the image to RGB and calculates skin-tone pixel color ratios (R > 95, G > 40, B > 20 with color variance checks). If skin-tone ratio > 0.20, confidence is heavily penalized, detecting face/portrait photos."),

        ("Q13: How does the system detect outdoor landscape photos?",
         "ProductDetectionService calculates sky blue pixel ratios (B > 140, B > R+20) and vegetation green pixel ratios (G > 100, G > R+25). If nature ratio > 0.25, confidence is penalized, preventing landscape photos from passing."),

        ("Q14: How does NirikshanX support Human-in-the-Loop enforcement?",
         "When AI confidence is low, evidence is missing, or image quality is borderline, the system sets overall_status to REVIEW_REQUIRED and generates a Case entity (e.g. NX-20260902-0010). Enforcement officers can review evidence, append notes, and override AI decisions before legal action."),

        ("Q15: How can NirikshanX scale for high-throughput production deployment?",
         "1. Decouple background AI worker tasks using Celery / Redis or AWS SQS instead of inline background tasks.\n"
         "2. Containerize with Docker & Kubernetes, running FastAPI behind NGINX load balancer.\n"
         "3. Migrate database from SQLite to PostgreSQL / AWS RDS.\n"
         "4. Deploy RapidOCR ONNX models on GPU/Triton Inference Server for batch image processing.")
    ]

    for q, a in qas:
        add_heading_2(doc, q)
        add_body_paragraph(doc, a)

    # Save output docx
    output_path = r"d:\NirikshanX\NirikshanX_Complete_Project_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation generated successfully at: {output_path}")

    # Also save to brain artifacts dir
    artifacts_dir = r"C:\Users\User\.gemini\antigravity-ide\brain\7852aa37-addf-4b87-95f5-4a6dd7e0a407"
    if os.path.exists(artifacts_dir):
        art_path = os.path.join(artifacts_dir, "NirikshanX_Complete_Project_Documentation.docx")
        doc.save(art_path)
        print(f"Documentation saved to brain artifacts at: {art_path}")

if __name__ == "__main__":
    create_project_documentation()
